import asyncio
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from fastmcp import FastMCP

mcp = FastMCP("Word Loader")

current_document_path = None

@mcp.tool()
async def set_document(document_name: str, document_path: str = None) -> str:
    """
    Set the current working document by name or path.
    
    Parameters:
        document_name (str): Name of the document (with or without .docx extension)
        document_path (str, optional): Full path to the document. If not provided, uses Desktop.
    Returns:
        str: Message indicating success or failure
    """
    global current_document_path
    
    try:
        if not document_name.endswith('.docx'):
            document_name += '.docx'
        
        if document_path:
            full_path = Path(document_path) / document_name
        else:
            desktop_path = Path.home() / "Desktop"
            full_path = desktop_path / document_name
        
        full_path.parent.mkdir(parents=True, exist_ok=True)
        
        if not full_path.exists():
            doc = Document()
            doc.save(str(full_path))
            message = f"Novo documento criado: {full_path.name}"
        else:
            message = f"Documento encontrado: {full_path.name}"
        
        current_document_path = str(full_path)
        return message
        
    except Exception as e:
        return f"Erro ao definir documento: {str(e)}"

@mcp.tool()
async def load_content(content: str) -> str:
    """
    Load content into the current Word document with markdown formatting.
    
    Parameters:
        content (str): The content to be loaded into the Word document (supports markdown).
    Returns:
        str: A message indicating the content has been loaded successfully.
    """
    global current_document_path
    
    if not current_document_path:
        desktop_path = Path.home() / "Desktop"
        current_document_path = str(desktop_path / "notas_claude.docx")
        
        if not Path(current_document_path).exists():
            doc = Document()
            doc.save(current_document_path)
    
    try:
        document = Document(current_document_path)
        
        if len(document.paragraphs) > 0 and document.paragraphs[-1].text.strip():
            document.add_paragraph("─" * 50)
        
        _process_markdown_content(document, content)
        
        document.save(current_document_path)
        
        doc_name = Path(current_document_path).name
        return f"Conteúdo salvo com sucesso em {doc_name}"
        
    except Exception as e:
        return f"Erro ao salvar conteúdo: {str(e)}"

def _process_markdown_content(document, content):
    """
    Process markdown content and add it to the Word document with proper formatting.
    """
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        if not line:
            i += 1
            continue
        
        if line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        
        elif line.startswith('- ') or line.startswith('* '):
            bullet_items = []
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ')):
                bullet_items.append(lines[i][2:])
                i += 1
            
            for item in bullet_items:
                p = document.add_paragraph(style='List Bullet')
                _add_formatted_text(p, item)
            continue
        
        elif re.match(r'^\d+\. ', line):
            numbered_items = []
            while i < len(lines) and re.match(r'^\d+\. ', lines[i]):
                numbered_items.append(re.sub(r'^\d+\. ', '', lines[i]))
                i += 1
            
            for item in numbered_items:
                p = document.add_paragraph(style='List Number')
                _add_formatted_text(p, item)
            continue
        
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            code_text = '\n'.join(code_lines)
            p = document.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            p.style = 'No Spacing'
        
        else:
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not _is_special_line(lines[i]):
                para_lines.append(lines[i].rstrip())
                i += 1
            
            para_text = ' '.join(para_lines)
            p = document.add_paragraph()
            _add_formatted_text(p, para_text)
            continue
        
        i += 1

def _is_special_line(line):
    """
    Check if a line is a special markdown element (header, list, etc.)
    """
    return (line.startswith('#') or 
            line.startswith('- ') or 
            line.startswith('* ') or 
            re.match(r'^\d+\. ', line) or
            line.startswith('```'))

def _add_formatted_text(paragraph, text):
    """
    Add text with markdown formatting to a paragraph.
    Supports **bold**, *italic*, `code`, and regular text.
    """
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    asyncio.run(mcp.run(transport="stdio"))